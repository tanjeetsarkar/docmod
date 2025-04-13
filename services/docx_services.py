# services/docx_service.py
import os
import uuid
from datetime import datetime
from typing import Any, Dict

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from sqlalchemy.orm import Session

from models.artifacts import Artifact


async def generate_docx_document(document_data: Dict[str, Any], db: Session) -> str:
    """
    Generate a DOCX document from document data
    
    Args:
        document_data: Dictionary containing document title, sections, etc.
        db: Database session
        
    Returns:
        Path to the generated DOCX file
    """
    # Create new Document
    doc = Document()
    
    # Add document title
    title = doc.add_heading(document_data["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add table of contents if needed
    toc_bookmark = None
    for section in document_data["sections"]:
        if section["type"] == "tableOfContents":
            toc_bookmark = add_toc(doc)
            doc.add_paragraph()  # Add space after TOC
            break
    
    # Process each section
    for section in document_data["sections"]:
        if section["type"] == "tableOfContents":
            # Already handled above
            continue
        elif section["type"] == "paragraph":
            add_paragraph_section(doc, section)
        elif section["type"] == "table":
            await add_table_section(doc, section, db)
        elif section["type"] == "image":
            await add_image_section(doc, section, db)
        elif section["type"] == "attachment":
            await add_attachment_section(doc, section, db)
    
    # Create a temporary file for the document
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    file_name = f"{document_data['title'].replace(' ', '_')}_{timestamp}.docx"
    output_dir = "generated_docs"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, file_name)
    
    # Save the document
    doc.save(file_path)
    
    # Update TOC if needed
    if toc_bookmark:
        update_toc(file_path)
    
    return file_path

def add_paragraph_section(doc: Document, section: Dict[str, Any]) -> None:
    """Add a paragraph section to the document."""
    # Add heading based on first line of text
    text = section["content"].strip()
    first_line = text.split('\n')[0]
    
    # Use the first line as a heading if it's relatively short
    if len(first_line) <= 100:
        heading = doc.add_heading(first_line, level=1)
        # Add bookmark for TOC
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(uuid.uuid4())[:8])
        bookmark_start.set(qn('w:name'), f"heading_{section['id']}")
        heading._element.append(bookmark_start)
        
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_start.get(qn('w:id')))
        heading._element.append(bookmark_end)
        
        # Add the rest of the paragraph if there's more text
        if len(text.split('\n')) > 1:
            remaining_text = '\n'.join(text.split('\n')[1:])
            p = doc.add_paragraph(remaining_text)
    else:
        # Add the entire text as a regular paragraph
        p = doc.add_paragraph(text)
    
    doc.add_paragraph()  # Add space after paragraph

async def add_table_section(doc: Document, section: Dict[str, Any], db: Session) -> None:
    """Add a table section to the document."""
    # Add heading for table
    heading = doc.add_heading(section["title"], level=2)
    
    # Fetch the artifact data
    artifact = db.query(Artifact).filter(Artifact.id == section["artifactId"]).first()
    if not artifact:
        doc.add_paragraph("Table data not found")
        return
    
    if artifact.data:
        # Create table from artifact data
        table_data = artifact.data
        
        # Determine number of rows and columns
        if isinstance(table_data, dict) and "data" in table_data:
            # Handle case where data is wrapped in a "data" key
            rows = table_data.get("data", [])
        elif isinstance(table_data, list):
            # Handle case where data is a list directly
            rows = table_data
        else:
            # Fallback for unknown format
            doc.add_paragraph("Invalid table data format")
            return
        
        if not rows:
            doc.add_paragraph("Table is empty")
            return
        
        # Get column headers
        if isinstance(rows[0], dict):
            # Data is a list of dictionaries
            headers = list(rows[0].keys())
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = str(header)
            
            # Add data rows
            for row_idx, row_data in enumerate(rows):
                cells = table.rows[row_idx + 1].cells
                for col_idx, header in enumerate(headers):
                    cells[col_idx].text = str(row_data.get(header, ""))
        elif isinstance(rows[0], list):
            # Data is a list of lists
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            
            # Add all rows
            for row_idx, row_data in enumerate(rows):
                cells = table.rows[row_idx].cells
                for col_idx, cell_value in enumerate(row_data):
                    cells[col_idx].text = str(cell_value)
    elif artifact.file_path:
        # Load data from CSV file
        try:
            df = pd.read_csv(artifact.file_path)
            headers = df.columns.tolist()
            
            table = doc.add_table(rows=len(df) + 1, cols=len(headers))
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = str(header)
            
            # Add data rows
            for row_idx, row_data in df.iterrows():
                cells = table.rows[row_idx + 1].cells
                for col_idx, header in enumerate(headers):
                    cells[col_idx].text = str(row_data[header])
        except Exception as e:
            doc.add_paragraph(f"Error loading table data: {str(e)}")
            return
    else:
        doc.add_paragraph("No table data available")
        return
        
    # Apply table style
    table.style = 'Table Grid'
    table.autofit = True
    
    doc.add_paragraph()  # Add space after table

async def add_image_section(doc: Document, section: Dict[str, Any], db: Session) -> None:
    """Add an image section to the document."""
    # Add heading for image
    heading = doc.add_heading(section["title"], level=2)
    
    # Fetch the artifact data
    artifact = db.query(Artifact).filter(Artifact.id == section["artifactId"]).first()
    if not artifact:
        doc.add_paragraph("Image not found")
        return
    
    if artifact.file_path and os.path.exists(artifact.file_path):
        # Add image to document
        try:
            doc.add_picture(artifact.file_path, width=Inches(6))
            
            # Add caption
            caption = doc.add_paragraph(f"Figure: {section['title']}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"Error adding image: {str(e)}")
    else:
        doc.add_paragraph("Image file not found")
    
    doc.add_paragraph()  # Add space after image

async def add_attachment_section(doc: Document, section: Dict[str, Any], db: Session) -> None:
    """Add an attachment reference section to the document."""
    # Add heading for attachment
    heading = doc.add_heading(f"Attachment: {section['title']}", level=2)
    
    # Fetch the artifact data
    artifact = db.query(Artifact).filter(Artifact.id == section["artifactId"]).first()
    if not artifact:
        doc.add_paragraph("Attachment not found")
        return
    
    if artifact.file_path and os.path.exists(artifact.file_path):
        # Add reference to attachment
        p = doc.add_paragraph(f"File: {os.path.basename(artifact.file_path)}")
        
        # You could potentially add hyperlink or embed the file here
        # depending on requirements
    else:
        doc.add_paragraph("Attachment file not found")
    
    doc.add_paragraph()  # Add space after attachment reference

def add_toc(doc: Document) -> str:
    """
    Add table of contents to document
    
    Returns:
        Bookmark name for TOC
    """
    heading = doc.add_heading("Table of Contents", level=1)
    
    # Create bookmark
    bookmark_name = "toc"
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_id = str(uuid.uuid4())[:8]
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    
    # Add placeholder paragraph for TOC
    p = doc.add_paragraph()
    p._element.append(bookmark_start)
    p._element.append(bookmark_end)
    
    # Add TOC field
    run = p.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld_char)
    
    instr_text = OxmlElement('w:instrText')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instr_text)
    
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_char)
    
    return bookmark_name

def update_toc(file_path: str) -> None:
    """
    Update table of contents in the document
    
    In a real application, you might need to use a Word COM object
    or an external library to update the TOC. This is a placeholder.
    """
    # This is just a placeholder. In real implementation, 
    # you would need to use a Word COM object or a specialized library
    # to update the TOC after creating the document.
    pass